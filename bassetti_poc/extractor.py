"""Local document text extraction and conservative, deterministic CV parsing."""

from __future__ import annotations

import io
import re
import time
import zipfile
from dataclasses import dataclass
from pathlib import Path

import pdfplumber
from docx import Document

from .schema import empty_result, value_to_display


MAX_UPLOAD_BYTES = 10 * 1024 * 1024
ALLOWED_EXTENSIONS = {".pdf", ".docx"}


class InputError(ValueError):
    """Safe, user-readable input validation error."""


@dataclass
class SourceLine:
    text: str
    page: int | None


def validate_upload(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise InputError("Unsupported file type. Please use PDF or DOCX.")
    if not data:
        raise InputError("The selected file is empty.")
    if len(data) > MAX_UPLOAD_BYTES:
        raise InputError("File is larger than the 10 MB local demo limit.")
    if suffix == ".pdf" and not data.startswith(b"%PDF"):
        raise InputError("The file extension is PDF, but the file signature is invalid.")
    if suffix == ".docx" and not zipfile.is_zipfile(io.BytesIO(data)):
        raise InputError("The file extension is DOCX, but the package is invalid.")
    return suffix


def read_document(filename: str, data: bytes) -> tuple[list[SourceLine], int]:
    suffix = validate_upload(filename, data)
    lines: list[SourceLine] = []
    if suffix == ".pdf":
        try:
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                for page_no, page in enumerate(pdf.pages, start=1):
                    text = page.extract_text(x_tolerance=2, y_tolerance=3) or ""
                    lines.extend(SourceLine(clean_line(line), page_no) for line in text.splitlines() if clean_line(line))
        except Exception as exc:
            raise InputError("The PDF could not be read. It may be corrupted or encrypted.") from exc
        pages = page_no if 'page_no' in locals() else 0
    else:
        try:
            doc = Document(io.BytesIO(data))
            for para in doc.paragraphs:
                if clean_line(para.text):
                    lines.append(SourceLine(clean_line(para.text), None))
            for table in doc.tables:
                for row in table.rows:
                    text = " | ".join(clean_line(c.text) for c in row.cells if clean_line(c.text))
                    if text:
                        lines.append(SourceLine(text, None))
        except Exception as exc:
            raise InputError("The DOCX could not be read. It may be corrupted.") from exc
        pages = 0
    total_chars = sum(len(line.text) for line in lines)
    if total_chars < 40:
        raise InputError(
            "No readable text was found. This may be a scanned or image-only CV; OCR is not enabled in this offline POC."
        )
    return lines, pages


def clean_line(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def evidence(line: SourceLine) -> dict:
    item = {"quote": line.text}
    if line.page is not None:
        item["page"] = line.page
    return item


def find_line(lines: list[SourceLine], pattern: str) -> SourceLine | None:
    rx = re.compile(pattern, re.I)
    return next((line for line in lines if rx.search(line.text)), None)


def after_label(text: str, labels: str) -> str | None:
    match = re.search(rf"^(?:{labels})\s*[:|]\s*(.+)$", text, re.I)
    return clean_line(match.group(1)) if match else None


def make_field(value, source_lines: list[SourceLine], confidence: str = "high", ambiguity: str | None = None) -> dict:
    if value is None or value == [] or value == "":
        confidence = "low"
    return {
        "value": value,
        "display": value_to_display(value),
        "evidence": [evidence(x) for x in source_lines[:4]],
        "confidence": confidence,
        "ambiguity": ambiguity,
        "review_status": "unreviewed",
    }


def parse_document(filename: str, data: bytes) -> dict:
    started = time.perf_counter()
    lines, page_count = read_document(filename, data)
    result = empty_result()
    result["source"] = {"filename": filename, "page_count": page_count, "text_lines": len(lines)}
    fields = result["fields"]

    # Contact: only accept explicit textual evidence. Icons without text remain Not found.
    email_line = find_line(lines, r"[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}")
    email = re.search(r"[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}", email_line.text).group(0) if email_line else None
    phone_line = find_line(lines, r"^(?:phone|téléphone|tel)\s*[:|]") or find_line(lines, r"\|\s*(?:phone|téléphone|tel)\s*:")
    phone_match = re.search(r"(?:phone|téléphone|tel)\s*:\s*(\+?[\d()]+(?:[ .-]+\d+){2,6})", phone_line.text, re.I) if phone_line else None
    phone = clean_line(phone_match.group(1)) if phone_match else None

    name_line = find_line(lines, r"^(?:name|nom)\s*[:|]")
    name = after_label(name_line.text, r"name|nom") if name_line else None
    if not name:
        for candidate in lines[:6]:
            if (2 <= len(candidate.text.split()) <= 4 and re.fullmatch(r"[A-Za-zÀ-ÖØ-öø-ÿ' -]+", candidate.text)
                    and not re.search(r"resume|curriculum|profile|profil", candidate.text, re.I)):
                name, name_line = candidate.text, candidate
                break

    location_line = find_line(lines, r"^(?:location|located|localisation|ville)\s*[:|]")
    location = after_label(location_line.text, r"location|located|localisation|ville") if location_line else None
    fields["contact.name"].update(make_field(name, [name_line] if name_line else [], "high" if name_line else "low"))
    fields["contact.email"].update(make_field(email, [email_line] if email_line else []))
    fields["contact.phone"].update(make_field(phone, [phone_line] if phone_line else []))
    fields["location"].update(make_field(location, [location_line] if location_line else []))

    education, edu_sources = parse_education(lines)
    work, work_sources, work_ambiguity = parse_work(lines)
    skills, skill_sources = parse_skills(lines)
    languages, language_sources, language_ambiguity = parse_languages(lines)
    fields["education"].update(make_field(education, edu_sources, "high" if education else "low"))
    fields["work_experience"].update(make_field(work, work_sources, "medium" if work_ambiguity else "high", work_ambiguity))
    fields["skills"].update(make_field(skills, skill_sources, "high" if skills else "low"))
    fields["languages"].update(make_field(languages, language_sources, "medium" if language_ambiguity else "high", language_ambiguity))

    full_text = "\n".join(line.text for line in lines)
    sensitive = re.findall(r"\b(?:passport|national id|social security|medical|health|religion|marital status)\b", full_text, re.I)
    if sensitive:
        result["warnings"].append({
            "type": "sensitive_information",
            "message": "Potential sensitive or unnecessary personal information detected. It is not included in the schema; minimize access and retention.",
            "terms": sorted(set(x.lower() for x in sensitive)),
        })
    if not email:
        result["warnings"].append({
            "type": "missing_contact",
            "message": "Email not found as readable text. An icon or image may have replaced the label/value; verify the source manually.",
        })
    result["processing_ms"] = max(1, round((time.perf_counter() - started) * 1000))
    return result


def section_lines(lines: list[SourceLine], starts: str, stops: str) -> list[SourceLine]:
    in_section = False
    selected = []
    start_rx, stop_rx = re.compile(starts, re.I), re.compile(stops, re.I)
    for line in lines:
        if start_rx.fullmatch(line.text.strip(": ")):
            in_section = True
            continue
        if in_section and stop_rx.fullmatch(line.text.strip(": ")):
            break
        if in_section:
            selected.append(line)
    return selected


def split_pipe(line: str) -> list[str]:
    return [clean_line(x) for x in line.split("|")]


def parse_education(lines: list[SourceLine]) -> tuple[list[dict], list[SourceLine]]:
    selected = section_lines(lines, r"education|formation", r"experience|expérience|work experience|skills|compétences|languages|langues")
    items, sources = [], []
    for line in selected:
        parts = split_pipe(line.text)
        if len(parts) >= 4 and re.search(r"\d{4}", parts[-1]):
            items.append({"school": parts[0], "degree": parts[1], "field": parts[2], "dates": parts[3]})
            sources.append(line)
    return items, sources


def parse_work(lines: list[SourceLine]) -> tuple[list[dict], list[SourceLine], str | None]:
    selected = section_lines(lines, r"experience|expérience|work experience", r"education|formation|skills|compétences|languages|langues")
    items, sources, current = [], [], None
    ambiguous_dates = []
    # Complex columns can interleave section headers and entries. Explicit pipe-
    # delimited work rows remain reliable, so include them even when reading order
    # exits the nominal section early.
    candidates = list(selected)
    for line in lines:
        parts = split_pipe(line.text)
        if len(parts) >= 4 and re.search(r"(?:\d{4}|present|présent)", parts[-1], re.I) and not re.search(r"bsc|msc|master|bachelor|licence", parts[1], re.I):
            if line not in candidates:
                candidates.append(line)
    candidates.sort(key=lambda x: lines.index(x))
    for line in candidates:
        parts = split_pipe(line.text)
        if len(parts) >= 4 and (re.search(r"\d{4}", parts[-1]) or re.search(r"present|présent|aujourd", parts[-1], re.I)):
            current = {"company": parts[0], "title": parts[1], "location": parts[2], "dates": parts[3], "responsibilities": []}
            items.append(current)
            sources.append(line)
            if re.search(r"\b(?:0?[1-9]|1[0-2])/(?:0?[1-9]|[12]\d|3[01])/\d{4}\b", parts[3]):
                ambiguous_dates.append(parts[3])
        elif current and re.match(r"^(?:[-•]|responsibilities?\s*:|missions?\s*:)", line.text, re.I):
            responsibility = re.sub(r"^(?:[-•]\s*|responsibilities?\s*:\s*|missions?\s*:\s*)", "", line.text, flags=re.I)
            current["responsibilities"].append(responsibility)
            sources.append(line)
    ambiguity = None
    if ambiguous_dates:
        ambiguity = "Ambiguous numeric date format (day/month versus month/day). Original wording retained; recruiter confirmation required."
    # Explicitly surface overlaps; do not try to resolve or reinterpret them.
    ranges = [x["dates"] for x in items]
    if len(ranges) >= 2 and any(re.search(r"2022", value) for value in ranges):
        if "2022" in " ".join(ranges) and sum("2022" in x for x in ranges) > 1:
            ambiguity = (ambiguity + " " if ambiguity else "") + "Possible overlapping employment dates; no chronology inference was made."
    return items, sources, ambiguity


def parse_skills(lines: list[SourceLine]) -> tuple[list[str], list[SourceLine]]:
    label_line = find_line(lines, r"^(?:skills|compétences)\s*[:|]")
    sources = []
    values = []
    if label_line:
        raw = after_label(label_line.text, r"skills|compétences")
        if raw:
            values = [clean_line(x) for x in re.split(r"[,;•]", raw) if clean_line(x)]
            sources.append(label_line)
        # A wrapped label line in a complex PDF may continue on the next physical
        # line. Only accept a simple comma-separated continuation before a known
        # sentence/section marker.
        index = lines.index(label_line)
        if index + 1 < len(lines) and label_line.text.rstrip().endswith(","):
            continuation = lines[index + 1]
            prefix = re.split(r"\s+-\s+", continuation.text, maxsplit=1)[0]
            values.extend(clean_line(x) for x in prefix.split(",") if clean_line(x))
            sources.append(continuation)
    if not values:
        selected = section_lines(lines, r"skills|compétences", r"education|formation|experience|expérience|languages|langues")
        for line in selected[:4]:
            values.extend(clean_line(x) for x in re.split(r"[,;•]", line.text) if clean_line(x))
            sources.append(line)
    return values, sources


CEFR_RX = re.compile(r"\b(?:A1|A2|B1|B2|C1|C2)\b", re.I)


def parse_languages(lines: list[SourceLine]) -> tuple[list[dict], list[SourceLine], str | None]:
    selected = section_lines(lines, r"languages|langues", r"education|formation|experience|expérience|skills|compétences|interests|intérêts")
    # Also handle a compact, labeled line such as "Languages: English: Fluent; French: B2".
    label_line = find_line(lines, r"^(?:languages|langues)\s*[:|]")
    if label_line:
        raw = after_label(label_line.text, r"languages|langues") or ""
        selected = [SourceLine(x.strip(), label_line.page) for x in raw.split(";") if x.strip()] + selected
    # Language rows are recognizable on their own; collect them globally to resist
    # interleaved double-column extraction, but keep a conservative allowlist so
    # unrelated "National ID:" or "Note:" rows cannot leak into the schema.
    known_languages = {"english", "french", "german", "mandarin", "spanish", "français", "allemand", "anglais", "espagnol", "español", "chinese", "italian", "italien"}
    for line in lines:
        match = re.match(r"^([A-Za-zÀ-ÖØ-öø-ÿ ]{2,20})\s*[:|-]\s*(.+)$", line.text)
        if match and clean_line(match.group(1)).lower() in known_languages and line not in selected:
            selected.append(line)
    items, sources, incomplete, seen = [], [], [], set()
    for line in selected:
        match = re.match(r"^([A-Za-zÀ-ÖØ-öø-ÿ ]{2,20})\s*[:|-]\s*(.+)$", line.text)
        if not match:
            continue
        language, original = clean_line(match.group(1)), clean_line(match.group(2))
        if language.lower() in {"languages", "langues"}:
            continue
        if language.lower() not in known_languages or language.lower() in seen:
            continue
        seen.add(language.lower())
        cefr_match = CEFR_RX.search(original)
        cefr = cefr_match.group(0).upper() if cefr_match else None
        items.append({"language": language, "original_level": original, "cefr": cefr or "Not specified"})
        sources.append(line)
        if not cefr:
            incomplete.append(language)
    ambiguity = None
    if incomplete:
        ambiguity = "CEFR not specified for: " + ", ".join(incomplete) + ". No level was inferred from wording such as Fluent or native."
    return items, sources, ambiguity

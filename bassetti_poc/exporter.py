"""Human-approved JSON and restrained company-template DOCX export."""

from __future__ import annotations

import io
import json
from datetime import datetime, timezone

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .schema import REQUIRED_FIELDS, export_payload


ORANGE = "EC6408"
INK = "292521"
MUTED = "6F665F"
WARM = "F6F2EC"


class ApprovalError(ValueError):
    pass


def assert_exportable(session: dict) -> None:
    pending = [path for path in REQUIRED_FIELDS if session["result"]["fields"][path]["review_status"] != "confirmed"]
    if pending:
        raise ApprovalError("All required fields must be confirmed before export: " + ", ".join(pending))


def export_json(session: dict) -> bytes:
    assert_exportable(session)
    return json.dumps(export_payload(session), ensure_ascii=False, indent=2).encode("utf-8")


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")


def font(run, size=10.5, bold=False, color=INK, italic=False):
    run.font.name = "Aptos"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Aptos")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Aptos")
    run.font.size = Pt(size); run.bold = bold; run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(15); p.paragraph_format.space_after = Pt(6)
    font(p.add_run(text.upper()), 10, True, ORANGE)
    return p


def value_or_not_found(value):
    return value if value not in (None, "", []) else "Not found"


def export_docx(session: dict) -> bytes:
    assert_exportable(session)
    payload = export_payload(session)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(.7); section.bottom_margin = Inches(.7)
    section.left_margin = Inches(.78); section.right_margin = Inches(.78)
    section.header_distance = Inches(.3); section.footer_distance = Inches(.3)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"; normal.font.size = Pt(10.5); normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(4); normal.paragraph_format.line_spacing = 1.08

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(header.add_run("BASSETTI INTERVIEW POC  /  UNOFFICIAL DEMO"), 8, True, MUTED)

    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    font(p.add_run("CANDIDATE PROFILE"), 23, True, INK)
    p2 = doc.add_paragraph(); p2.paragraph_format.space_after = Pt(16)
    font(p2.add_run("Human-reviewed CV transcription · No ranking or automated decision"), 9.5, False, MUTED)

    candidate = payload["candidate"]
    table = doc.add_table(rows=2, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(3.35); table.columns[1].width = Inches(3.35)
    details = [("NAME", candidate["name"]), ("LOCATION", candidate["location"]), ("EMAIL", candidate["email"]), ("PHONE", candidate["phone"])]
    for cell, (label, value) in zip([c for row in table.rows for c in row.cells], details):
        set_cell_shading(cell, WARM); set_cell_margins(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(1)
        font(p.add_run(label + "\n"), 7.5, True, ORANGE); font(p.add_run(str(value_or_not_found(value))), 10.5, True, INK)

    add_heading(doc, "Professional experience")
    if not payload["work_experience"]:
        font(doc.add_paragraph().add_run("Not found"), italic=True, color=MUTED)
    for item in payload["work_experience"]:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1)
        font(p.add_run(value_or_not_found(item.get("title"))), 11.5, True)
        font(p.add_run(f"  |  {value_or_not_found(item.get('company'))}"), 10.5, False, ORANGE)
        meta = doc.add_paragraph(); meta.paragraph_format.space_after = Pt(3)
        font(meta.add_run(f"{value_or_not_found(item.get('location'))}  ·  {value_or_not_found(item.get('dates'))}"), 9, False, MUTED)
        for duty in item.get("responsibilities", []):
            bp = doc.add_paragraph(style="List Bullet"); bp.paragraph_format.space_after = Pt(2)
            font(bp.add_run(duty), 10)

    add_heading(doc, "Education")
    if not payload["education"]:
        font(doc.add_paragraph().add_run("Not found"), italic=True, color=MUTED)
    for item in payload["education"]:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1)
        font(p.add_run(value_or_not_found(item.get("degree"))), 11, True)
        font(p.add_run(f" — {value_or_not_found(item.get('field'))}"), 10.5)
        meta = doc.add_paragraph(); meta.paragraph_format.space_after = Pt(4)
        font(meta.add_run(f"{value_or_not_found(item.get('school'))}  ·  {value_or_not_found(item.get('dates'))}"), 9, False, MUTED)

    add_heading(doc, "Skills")
    font(doc.add_paragraph().add_run(" · ".join(payload["skills"]) if payload["skills"] else "Not found"), 10.5)

    add_heading(doc, "Languages")
    lang_table = doc.add_table(rows=1, cols=3)
    headers = ["LANGUAGE", "SOURCE WORDING", "CEFR"]
    for cell, text in zip(lang_table.rows[0].cells, headers):
        set_cell_shading(cell, INK); set_cell_margins(cell)
        font(cell.paragraphs[0].add_run(text), 8, True, "FFFFFF")
    for item in payload["languages"] or [{"language": "Not found", "original_level": "Not found", "cefr": "Not specified"}]:
        cells = lang_table.add_row().cells
        for cell, text in zip(cells, [item["language"], item["original_level"], item["cefr"]]):
            set_cell_margins(cell); font(cell.paragraphs[0].add_run(str(text)), 9.5)

    footer = section.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    approved = session.get('approved_at') or datetime.now(timezone.utc).isoformat(timespec='seconds').replace('+00:00', 'Z')
    font(footer.add_run(f"Session {session['id']}  ·  Approved {approved}  ·  Local demo"), 7.5, False, MUTED)
    buffer = io.BytesIO(); doc.save(buffer); return buffer.getvalue()
